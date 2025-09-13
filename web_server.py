from flask import Flask, request, send_file, Response, jsonify
from werkzeug.utils import secure_filename
import os
import json
from advanced_docx_translator import translate_docx_advanced, translation_cache #, summarize_text_groq
# from pdf_to_word import convert_pdf_exactly
# from pdf_to_word import smart_pdf_to_word
from pdf_to_word import convert_pdf_to_docx_pdf2docx, convert_pdf_with_aspose
import logging
import time
import threading
from docx import Document
import uuid
import shutil
import io # Import io for in-memory file handling
import filecmp

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Optionally load environment variables from a .env file if present
try:
    from dotenv import load_dotenv  # type: ignore
    # Load .env explicitly from this file's directory to avoid CWD issues
    _DOTENV_PATH = os.path.join(os.path.dirname(__file__), '.env')
    load_dotenv(dotenv_path=_DOTENV_PATH)
    logger.info(f"Environment variables loaded from {_DOTENV_PATH} (if present).")
    # Log detected keys (masking values)
    _groq = bool(os.environ.get('GROQ_API_KEY') or os.environ.get('GROQ_KEY'))
    _gem = bool(os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY'))
    logger.info(f"API key presence — Groq: {_groq}, Gemini: {_gem}")
except Exception as _e:
    logger.warning(f"Could not load .env file: {_e}")

class TranslationProgress:
    def __init__(self):
        self.progress = 0
        self.status = "Ready"
        self.api_calls = 0
        self.cached = 0
        self.messages = []
        self.summary = None

    def update(self, progress=None, status=None, api_calls=None, cached=None, message=None, summary=None):
        if progress is not None:
            self.progress = progress
        if status is not None:
            self.status = status
        if api_calls is not None:
            self.api_calls = api_calls
        if cached is not None:
            self.cached = cached
        if message is not None:
            self.messages.append(message)
            logger.info(message)
        if summary is not None:
            self.summary = summary

    def to_json(self):
        data = {
            'progress': self.progress,
            'status': self.status,
            'apiCalls': self.api_calls,
            'cached': self.cached,
            'message': self.messages[-1] if self.messages else None
        }
        if self.summary is not None:
            data['summary'] = self.summary
        return json.dumps(data)

    def to_final_json(self, download_url=None):
        """Generate JSON for the final update, optionally including download URL."""
        data = {
            'progress': 100,
            'status': self.status,
            'apiCalls': self.api_calls,
            'cached': self.cached,
            'message': self.messages[-1] if self.messages else None
        }
        if self.summary is not None:
            data['summary'] = self.summary
        if download_url:
            data['downloadUrl'] = download_url
        return json.dumps(data)

@app.route('/')
def index():
    return send_file('web_translator.html')

@app.route('/styles.css')
def styles():
    return send_file('styles.css')

@app.route('/translator.js')
def translator_js():
    return send_file('translator.js')

@app.route('/translate', methods=['POST'])
def translate():
    """
    Main translation endpoint. Handles file upload, optional PDF conversion, translation, summary, and download.
    Modularized for clarity and maintainability.
    """
    def error_response(message, code=400):
        logger.error(message)
        return message, code

    # --- 1. Validate and Save Uploaded File ---
    if 'file' not in request.files:
        return error_response('No file provided')
    file = request.files['file']
    if file.filename == '':
        return error_response('No file selected')
    if not (file.filename.lower().endswith('.docx') or file.filename.lower().endswith('.pdf')):
        return error_response('Invalid file type. Please upload a DOCX or PDF file.')
    original_filename = secure_filename(file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], original_filename)
    file.save(input_path)
    logger.info(f"File saved to {input_path}")

    # --- 2. Parse Options ---
    target_lang = request.form.get('targetLang', 'es')
    engine = request.form.get('engine', 'gemini')
    generate_summary = request.form.get('generateSummary') == 'true'
    first_page_only = request.form.get('firstPageOnly') == 'true'
    tone = request.form.get('tone', 'professional')
    pdf_engine = request.form.get('pdfEngine', 'pdf2docx')

    # --- 3. PDF to DOCX Conversion (if needed) ---
    def convert_pdf(pdf_path, pdf_engine, first_page_only):
        docx_filename = f"{os.path.splitext(original_filename)[0]}.docx"
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        
        try:
            if first_page_only:
                start_page = 1  # Aspose uses 1-based indexing
                end_page = 1
            else:
                start_page = None
                end_page = None

            if pdf_engine == 'pdf2docx':
                # pdf2docx uses 0-based indexing
                start = 0 if first_page_only else None
                end = 1 if first_page_only else None
                convert_pdf_to_docx_pdf2docx(pdf_path, docx_path, start=start, end=end)
            elif pdf_engine == 'aspose':
                convert_pdf_with_aspose(pdf_path, docx_path, start=start_page, end=end_page)
            else:
                raise ValueError(f"Unsupported PDF conversion engine: {pdf_engine}")
                
            if not os.path.exists(docx_path):
                raise Exception("Conversion failed: Output file not created")
                
            return docx_path
            
        except Exception as e:
            logger.error(f"PDF conversion failed using {pdf_engine}: {str(e)}")
            raise

    docx_file_to_process = input_path
    current_filename = original_filename
    if original_filename.lower().endswith('.pdf'):
        try:
            docx_file_to_process = convert_pdf(input_path, pdf_engine, first_page_only)
            current_filename = os.path.basename(docx_file_to_process)
            # Remove original PDF after conversion
            if os.path.exists(input_path):
                os.remove(input_path)
        except Exception as e:
            return error_response(f"PDF conversion failed: {str(e)}", 500)

    # --- 4. Define Output Paths ---
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'translated_{os.path.splitext(current_filename)[0]}.docx')
    download_path = os.path.join(app.config['UPLOAD_FOLDER'], f'download_{os.path.splitext(current_filename)[0]}.docx')

    # --- 5. Streaming Response Generator ---
    def generate():
        progress = TranslationProgress()
        translation_complete = False
        translation_error = None
        summary_text = None
        translation_thread = None

        # --- 5a. Pre-calculate Total Paragraphs ---
        try:
            progress.update(status="Analyzing document structure...", progress=1)
            yield progress.to_json() + '\n'
            temp_doc = Document(docx_file_to_process)
            count = len(temp_doc.paragraphs)
            for table in temp_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        count += len(cell.paragraphs)
            for section in temp_doc.sections:
                if section.header: count += len(section.header.paragraphs)
                if section.footer: count += len(section.footer.paragraphs)
            total_paragraphs_for_progress = max(1, count)
            progress.update(message=f"Document has approx. {total_paragraphs_for_progress} paragraphs to process.")
            yield progress.to_json() + '\n'
            del temp_doc
        except Exception as e:
            total_paragraphs_for_progress = 1
            progress.update(message="Could not pre-calculate total paragraphs. Progress may be less accurate.")
            yield progress.to_json() + '\n'

        # --- 5b. Progress Callback ---
        def update_progress_callback(processed_count, total_count):
            base_progress = 10 if generate_summary else 5
            max_translation_progress = 95
            if total_count > 0:
                percent = base_progress + int(((processed_count / total_count) * (max_translation_progress - base_progress)))
            else:
                percent = base_progress
            percent = min(max_translation_progress, percent)
            progress.update(progress=percent, status=f"Translating ({processed_count}/{total_count})...")

        # --- 5c. Generate Summary (if requested) ---
        if generate_summary:
            progress.update(status="Generating summary...", progress=2, message="Reading document for summary...")
            yield progress.to_json() + '\n'
            try:
                summary_doc = Document(docx_file_to_process)
                full_text_list = []
                for para in summary_doc.paragraphs:
                    full_text_list.append(para.text)
                for table in summary_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full_text_list.append(para.text)
                full_text = "\n".join(full_text_list).strip()
                if full_text:
                    progress.update(status="Generating summary...", message="Sending text to Groq for summarization...")
                    yield progress.to_json() + '\n'
                    #summary_text = summarize_text_groq(full_text, target_language=target_lang, tone=tone)
                    progress.update(status="Summary generated", message="Summary received.", summary=summary_text)
                    yield progress.to_json() + '\n'
                else:
                    progress.update(status="Skipping summary", message="Document contains no text to summarize.")
                    yield progress.to_json() + '\n'
            except Exception as e:
                progress.update(status="Summary Error", message=f"Failed to generate summary: {e}")
                yield progress.to_json() + '\n'

        # --- 5d. Translation Thread ---
        progress.update(status="Starting translation...", progress=max(progress.progress, 5))
        yield progress.to_json() + '\n'
        def run_translation():
            nonlocal translation_complete, translation_error
            try:
                result = translate_docx_advanced(
                    docx_file_to_process,
                    output_path,
                    engine=engine.lower(),
                    target_language=target_lang,
                    progress_callback=update_progress_callback,
                    # first_page_only=first_page_only,
                    tone=tone
                )
                if result:
                    # Copy result to download path
                    shutil.copy2(output_path, download_path)
                    # Detect no-op translations: if output is byte-identical to input, treat as failure
                    try:
                        identical = False
                        if os.path.exists(docx_file_to_process) and os.path.exists(output_path):
                            identical = filecmp.cmp(docx_file_to_process, output_path, shallow=False)
                        if identical:
                            translation_complete = False
                            translation_error = "Translation produced an output identical to the input (no changes detected)."
                            logger.warning(f"Translation output identical to input: {docx_file_to_process} == {output_path}")
                        else:
                            translation_complete = True
                    except Exception as e:
                        # If comparison fails, still mark translation as complete but log the error
                        logger.warning(f"Could not compare input/output files for changes: {e}")
                        translation_complete = True
                else:
                    translation_complete = False
                    if not translation_error:
                        translation_error = "Translation was aborted or failed."
            except Exception as e:
                translation_error = str(e)
                translation_complete = False
        translation_thread = threading.Thread(target=run_translation)
        translation_thread.daemon = True
        translation_thread.start()

        # --- 5e. Monitor Progress ---
        last_yield_time = time.time()
        while translation_thread.is_alive():
            current_time = time.time()
            if current_time - last_yield_time >= 1.0:
                yield progress.to_json() + '\n'
                last_yield_time = current_time
            time.sleep(0.1)

        # --- 5f. Final Update ---
        final_download_url = None
        if translation_complete:
            final_message = "Document translation completed successfully!"
            download_base_name = os.path.splitext(current_filename)[0]
            final_download_url = f"/download/{download_base_name}"
            progress.update(status="Translation completed", message=final_message, api_calls=len(translation_cache))
        elif translation_error:
            progress.update(status="Error", message=f"Translation failed: {translation_error}", summary=None)
        else:
            progress.update(status="Unknown State", message="Translation process ended unexpectedly.")
        yield progress.to_final_json(download_url=final_download_url) + '\n'

        # --- 5g. Cleanup (optional, can be re-enabled) ---
        for f_path in [output_path, docx_file_to_process]:
            if f_path and os.path.exists(f_path) and f_path != download_path:
                try:
                    os.remove(f_path)
                except Exception as e:
                    logger.error(f"Error cleaning up file {f_path}: {str(e)}")

    return Response(generate(),
                    mimetype='text/event-stream',
                    headers={
                        'Cache-Control': 'no-cache',
                        'Connection': 'keep-alive',
                        'X-Accel-Buffering': 'no'
                    })

@app.route('/export-summary', methods=['POST'])
def export_summary():
    """Exports the provided summary text as a DOCX file."""
    logger.info("Received request to export summary.")
    try:
        data = request.get_json()
        summary_text = data.get('summary')

        if not summary_text:
            logger.warning("Export summary request received with no summary text.")
            return jsonify({'error': 'No summary text provided'}), 400

        # Create a new DOCX document in memory
        doc = Document()
        doc.add_paragraph(summary_text)

        # Save DOCX to a byte stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0) # Reset stream position to the beginning

        logger.info("Successfully created summary DOCX in memory.")

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='document_summary.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Error exporting summary: {e}", exc_info=True)
        return jsonify({'error': f'An internal error occurred: {e}'}), 500

@app.route('/clear-cache', methods=['POST'])
def clear_cache():
    count = len(translation_cache)
    translation_cache.clear()
    return jsonify({
        'message': f'Cleared {count} cached translations.',
        'status': 'success'
    })

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download the translated document"""
    logger.info(f"Download request received for filename from URL: {filename}")
    # Use the filename from the URL directly as the base_name
    base_name = filename 
    logger.info(f"Using base_name directly from URL: {base_name}")
    
    download_path = os.path.join(app.config['UPLOAD_FOLDER'], f'download_{base_name}.docx')
    logger.info(f"Constructed download path: {download_path}")
    
    # --- Add Directory Listing for Debugging --- 
    try:
        upload_dir_contents = os.listdir(app.config['UPLOAD_FOLDER'])
        logger.info(f"Contents of {app.config['UPLOAD_FOLDER']}: {upload_dir_contents}")
    except Exception as list_err:
        logger.error(f"Error listing directory {app.config['UPLOAD_FOLDER']}: {list_err}")
    # --- End Directory Listing --- 
    
    # Explicitly check existence and log result
    file_exists = os.path.exists(download_path)
    logger.info(f"Checking existence of {download_path}: {file_exists}")
    
    if not file_exists:
        logger.error(f"File not found at path: {download_path}")
        return "File not found or processing incomplete", 404
    
    # Set the download name (you might want to adjust this if original filenames had extensions)
    # For now, keeping the pattern consistent with potential original names
    download_name = f"translated_{base_name}.docx"
    logger.info(f"Serving file {download_path} as attachment name {download_name}")
    
    return send_file(download_path, 
                    as_attachment=True, 
                    download_name=download_name, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Scheduled task to clean up old files
def cleanup_old_files():
    """Clean up files older than 30 minutes"""
    while True:
        try:
            current_time = time.time()
            for file in os.listdir(app.config['UPLOAD_FOLDER']):
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getmtime(file_path)
                    # Remove files older than 30 minutes
                    if file_age > 30 * 60:
                        os.remove(file_path)
                        logger.info(f"Removed old file: {file_path}")
        except Exception as e:
            logger.error(f"Error cleaning up old files: {e}")
        
        # Run every 10 minutes
        time.sleep(10 * 60)

# Start the cleanup thread when the app starts
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

if __name__ == '__main__':
    app.run(debug=True, port=5000)


