import os
import io
import json
import time
import uuid
import shutil
import logging
import threading
import filecmp
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, Response, HTTPException
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse, PlainTextResponse
from fastapi.middleware.cors import CORSMiddleware

from docx import Document

from advanced_docx_translator import translate_docx_advanced, translation_cache
from pdf_to_word import convert_pdf_to_docx_pdf2docx


app = FastAPI()

# CORS (adjust origins in production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
MAX_UPLOAD_BYTES = 16 * 1024 * 1024
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TranslationProgress:
    def __init__(self):
        self.progress = 0
        self.status = "Ready"
        self.api_calls = 0
        self.cached = 0
        self.messages = []
        self.summary = None

    def update(self, progress=None, status=None, api_calls=None, cached=None, message=None, summary=None):
        if progress is not None:
            self.progress = progress
        if status is not None:
            self.status = status
        if api_calls is not None:
            self.api_calls = api_calls
        if cached is not None:
            self.cached = cached
        if message is not None:
            self.messages.append(message)
            logger.info(message)
        if summary is not None:
            self.summary = summary

    def to_json(self):
        data = {
            'progress': self.progress,
            'status': self.status,
            'apiCalls': self.api_calls,
            'cached': self.cached,
            'message': self.messages[-1] if self.messages else None
        }
        if self.summary is not None:
            data['summary'] = self.summary
        return json.dumps(data)

    def to_final_json(self, download_url=None):
        data = {
            'progress': 100,
            'status': self.status,
            'apiCalls': self.api_calls,
            'cached': self.cached,
            'message': self.messages[-1] if self.messages else None
        }
        if self.summary is not None:
            data['summary'] = self.summary
        if download_url:
            data['downloadUrl'] = download_url
        return json.dumps(data)


@app.get("/health")
def health():
    return {"status": "ok"}

@app.get("/")
def root():
    return {"status": "ok"}


@app.post("/translate")
async def translate(
    file: UploadFile = File(...),
    targetLang: str = Form('es'),
    engine: str = Form('gemini'),
    generateSummary: Optional[bool] = Form(False),
    firstPageOnly: Optional[bool] = Form(False),
    tone: str = Form('professional'),
    pdfEngine: str = Form('pdf2docx'),
    max_total_chars: Optional[str] = Form(None)

):
    # Validate file type and size
    if file is None or file.filename == "":
        raise HTTPException(status_code=400, detail="No file provided")
    lower = file.filename.lower()
    if not (lower.endswith('.docx') or lower.endswith('.pdf')):
        raise HTTPException(status_code=400, detail='Invalid file type. Please upload a DOCX or PDF file.')

    # Save upload to disk (bounded read)
    original_filename = os.path.basename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, original_filename)
    size = 0
    with open(input_path, 'wb') as out:
        while True:
            chunk = await file.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > MAX_UPLOAD_BYTES:
                out.close()
                try:
                    os.remove(input_path)
                except Exception:
                    pass
                raise HTTPException(status_code=413, detail='Uploaded file too large')
            out.write(chunk)
    logger.info(f"File saved to {input_path}")

    # PDF conversion if needed
    docx_file_to_process = input_path
    current_filename = original_filename
    if original_filename.lower().endswith('.pdf'):
        try:
            def convert_pdf(pdf_path, pdf_engine, first_page_only):
                docx_filename = f"{os.path.splitext(original_filename)[0]}.docx"
                docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)
                if first_page_only:
                    start_page = 1
                    end_page = 1
                else:
                    start_page = None
                    end_page = None
                if pdf_engine == 'pdf2docx':
                    start = 0 if first_page_only else None
                    end = 1 if first_page_only else None
                    convert_pdf_to_docx_pdf2docx(pdf_path, docx_path, start=start, end=end)
            
                else:
                    raise ValueError(f"Unsupported PDF conversion engine: {pdf_engine}")
                if not os.path.exists(docx_path):
                    raise RuntimeError("Conversion failed: Output file not created")
                return docx_path

            docx_file_to_process = convert_pdf(input_path, pdfEngine, firstPageOnly)
            current_filename = os.path.basename(docx_file_to_process)
            if os.path.exists(input_path):
                os.remove(input_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF conversion failed: {e}")

    # Output paths
    output_path = os.path.join(UPLOAD_FOLDER, f'translated_{os.path.splitext(current_filename)[0]}.docx')
    download_path = os.path.join(UPLOAD_FOLDER, f'download_{os.path.splitext(current_filename)[0]}.docx')

    # Streaming generator (SSE-friendly but sent as text/event-stream)
    def event_stream():
        progress = TranslationProgress()
        translation_complete = False
        translation_error = None

        # Pre-calc approx paragraphs
        try:
            progress.update(status="Analyzing document structure...", progress=1)
            yield progress.to_json() + "\n"
            temp_doc = Document(docx_file_to_process)
            count = len(temp_doc.paragraphs)
            for table in temp_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        count += len(cell.paragraphs)
            for section in temp_doc.sections:
                if section.header: count += len(section.header.paragraphs)
                if section.footer: count += len(section.footer.paragraphs)
            total_paragraphs_for_progress = max(1, count)
            progress.update(message=f"Document has approx. {total_paragraphs_for_progress} paragraphs to process.")
            yield progress.to_json() + "\n"
        except Exception:
            progress.update(message="Could not pre-calculate total paragraphs. Progress may be less accurate.")
            yield progress.to_json() + "\n"

        def update_progress_callback(processed_count, total_count):
            base_progress = 10 if generateSummary else 5
            max_translation_progress = 95
            if total_count > 0:
                percent = base_progress + int(((processed_count / total_count) * (max_translation_progress - base_progress)))
            else:
                percent = base_progress
            percent = min(max_translation_progress, percent)
            progress.update(progress=percent, status=f"Translating ({processed_count}/{total_count})...")

        def run_translation():
            nonlocal translation_complete, translation_error
           
           
      
            print(type(max_total_chars))
            try:
                result = translate_docx_advanced(
                    docx_file_to_process,
                    output_path,
                    engine=engine.lower(),
                    target_language=targetLang,
                    progress_callback=update_progress_callback,
                    tone=tone,
                    max_total_chars=5000
                )
                if result:
                    shutil.copy2(output_path, download_path)
                    try:
                        identical = False
                        if os.path.exists(docx_file_to_process) and os.path.exists(output_path):
                            identical = filecmp.cmp(docx_file_to_process, output_path, shallow=False)
                        if identical:
                            translation_complete = False
                            translation_error = "Translation produced an output identical to the input (no changes detected)."
                            logger.warning(f"Translation output identical to input: {docx_file_to_process} == {output_path}")
                        else:
                            translation_complete = True
                    except Exception as e:
                        logger.warning(f"Could not compare input/output files for changes: {e}")
                        translation_complete = True
                else:
                    translation_complete = False
                    if not translation_error:
                        translation_error = "Translation was aborted or failed."
            except Exception as e:
                translation_error = str(e)
                translation_complete = False

        t = threading.Thread(target=run_translation, daemon=True)
        t.start()

        last_yield_time = time.time()
        while t.is_alive():
            now = time.time()
            if now - last_yield_time >= 1.0:
                yield progress.to_json() + "\n"
                last_yield_time = now
            time.sleep(0.1)

        final_download_url = None
        if translation_complete:
            final_message = "Document translation completed successfully!"
            download_base_name = os.path.splitext(current_filename)[0]
            final_download_url = f"/download/{download_base_name}"
            progress.update(status="Translation completed", message=final_message, api_calls=len(translation_cache))
        elif translation_error:
            progress.update(status="Error", message=f"Translation failed: {translation_error}", summary=None)
        else:
            progress.update(status="Unknown State", message="Translation process ended unexpectedly.")
        yield progress.to_final_json(download_url=final_download_url) + "\n"

        # Optional cleanup
        for f_path in [output_path, docx_file_to_process]:
            if f_path and os.path.exists(f_path) and f_path != download_path:
                try:
                    os.remove(f_path)
                except Exception as e:
                    logger.error(f"Error cleaning up file {f_path}: {e}")

    return StreamingResponse(event_stream(), media_type='text/event-stream', headers={
        'Cache-Control': 'no-cache',
        'Connection': 'keep-alive',
        'X-Accel-Buffering': 'no'
    })


@app.get('/download/{filename}')
def download_file(filename: str):
    base_name = filename
    download_path = os.path.join(UPLOAD_FOLDER, f'download_{base_name}.docx')
    if not os.path.exists(download_path):
        raise HTTPException(status_code=404, detail='File not found or processing incomplete')
    return FileResponse(download_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f'translated_{base_name}.docx')


@app.post('/export-summary')
async def export_summary(data: dict):
    summary_text = data.get('summary') if data else None
    if not summary_text:
        return JSONResponse({'error': 'No summary text provided'}, status_code=400)
    doc = Document()
    doc.add_paragraph(summary_text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={
        'Content-Disposition': 'attachment; filename="document_summary.docx"'
    })


@app.post('/clear-cache')
def clear_cache():
    count = len(translation_cache)
    translation_cache.clear()
    return { 'message': f'Cleared {count} cached translations.', 'status': 'success' }


def cleanup_old_files():
    while True:
        try:
            current_time = time.time()
            for file_name in os.listdir(UPLOAD_FOLDER):
                file_path = os.path.join(UPLOAD_FOLDER, file_name)
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getmtime(file_path)
                    if file_age > 30 * 60:
                        os.remove(file_path)
                        logger.info(f"Removed old file: {file_path}")
        except Exception as e:
            logger.error(f"Error cleaning up old files: {e}")
        time.sleep(10 * 60)


# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()


# Run with: uvicorn main:app --host 0.0.0.0 --port 8000




